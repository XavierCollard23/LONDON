#!/usr/bin/env python3
"""
Optimise et reorganise l'itineraire "London - mise a jour.docx".

Fonctions principales :
  - Analyse le document Word existant et extrait les activites.
  - Identifie les doublons d'activites et les regroupe sur la journee la plus logique.
  - Ajoute des incontournables manquants en respectant les thematiques quotidiennes.
  - Calcule une sequence optimisee (proximite geographique + deplacements realistes).
  - Genere une carte Folium par jour ainsi qu'un nouveau document Word synthetique.
"""

from __future__ import annotations

import argparse
import json
import math
import os
import re
import unicodedata
from urllib.parse import quote
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from zipfile import ZipFile

import xml.etree.ElementTree as ET

from branca.element import MacroElement
from jinja2 import Template
try:
    import folium
except ImportError as exc:  # pragma: no cover - dependance exterieure
    raise SystemExit("Le module folium est requis. Installez-le avant dexecuter ce script.") from exc

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Le module python-docx est requis. Installez-le avant dexecuter ce script.") from exc


# ---------------------------------------------------------------------------
# Modele de donnees
# ---------------------------------------------------------------------------

DAY_MARKERS = ["\U0001F5D3", "\U0001F5D3\ufe0f"]

@dataclass
class LocationInfo:
    name: str
    lat: float
    lon: float
    category: str
    default_duration: int
    aliases: Sequence[str] = field(default_factory=list)
    notes: str = ""


@dataclass
class TimelineItem:
    time: str
    activity: str
    details: str


@dataclass
class DaySection:
    index: int
    title: str
    theme: str
    lines: List[str]
    timeline: List[TimelineItem]
    locations: List[str] = field(default_factory=list)
    removed_duplicates: List[str] = field(default_factory=list)
    added_essentials: List[str] = field(default_factory=list)


@dataclass
class Segment:
    start: int
    end: int
    title: str
    location: str
    details: str
    segment_type: str  # "visit" | "transfer" | "meal" | "buffer"


# ---------------------------------------------------------------------------
# Donnees geographiques et configuration
# ---------------------------------------------------------------------------


def _loc(
    name: str,
    lat: float,
    lon: float,
    category: str,
    duration: int,
    aliases: Sequence[str] = (),
    notes: str = "",
) -> LocationInfo:
    return LocationInfo(name=name, lat=lat, lon=lon, category=category, default_duration=duration, aliases=aliases, notes=notes)


LOCATION_LIBRARY: Sequence[LocationInfo] = [
    _loc("Heathrow Airport", 51.4700, -0.4543, "transport", 30, aliases=("heathrow", "arrivee heathrow")),
    _loc("Paddington Station", 51.5154, -0.1754, "transport", 15, aliases=("paddington",)),
    _loc("Southwark Station", 51.5046, -0.1050, "transport", 10, aliases=("southwark",)),
    _loc("citizenM London Bankside", 51.5055, -0.0993, "hotel", 30, aliases=("citizenm", "citizenm bankside", "citizenm london bankside")),
    _loc("South Bank Promenade", 51.5073, -0.0995, "walk", 60, aliases=("south bank",)),
    _loc("London Eye", 51.5033, -0.1196, "landmark", 60, aliases=("london eye",)),
    _loc("Carnaby Street", 51.5135, -0.1394, "walk", 45, aliases=("carnaby",)),
    _loc("Oxford Circus", 51.5150, -0.1410, "landmark", 10, aliases=("oxford circus",)),
    _loc("Regent Street", 51.5145, -0.1420, "walk", 30, aliases=("regent street",)),
    _loc("Oxford Street", 51.5154, -0.1410, "walk", 30, aliases=("oxford street",)),
    _loc("Covent Garden", 51.5129, -0.1226, "landmark", 45, aliases=("covent garden",)),
    _loc("Seven Dials Market", 51.5135, -0.1258, "food", 60, aliases=("seven dials",)),
    _loc("Soho", 51.5136, -0.1365, "walk", 60, aliases=("soho",)),
    _loc("Liberty London", 51.5136, -0.1406, "shopping", 45, aliases=("liberty london", "liberty"), notes="Grand magasin emblematique."),
    _loc("Mayfair Arcades", 51.5090, -0.1488, "walk", 60, aliases=("mayfair",)),
    _loc("Harrods", 51.4996, -0.1635, "shopping", 75, aliases=("harrods",)),
    _loc("New Bond Street", 51.5123, -0.1497, "shopping", 45, aliases=("bond street", "new bond street")),
    _loc("Sketch Mayfair", 51.5114, -0.1417, "food", 75, aliases=("sketch",)),
    _loc("Buckingham Palace", 51.5014, -0.1419, "landmark", 40, aliases=("buckingham", "buckingham palace")),
    _loc("St James's Park", 51.5010, -0.1365, "park", 40, aliases=("st james", "st james's park")),
    _loc("Hyde Park Winter Wonderland", 51.5073, -0.1668, "experience", 150, aliases=("winter wonderland", "hyde park winter wonderland")),
    _loc("Hyde Park", 51.5073, -0.1657, "park", 60, aliases=("hyde park",)),
    _loc("Bavarian Village", 51.5071, -0.1689, "food", 60, aliases=("bavarian village",)),
    _loc("Alpen Village", 51.5080, -0.1698, "food", 45, aliases=("alpen village",)),
    _loc("Bankside Pier", 51.5079, -0.0980, "transport", 10, aliases=("bankside pier",)),
    _loc("Westminster Pier", 51.5010, -0.1230, "transport", 10, aliases=("westminster pier",)),
    _loc("Palace of Westminster", 51.4995, -0.1245, "landmark", 30, aliases=("palais de westminster", "westminster", "big ben")),
    _loc("Buckingham Gate", 51.4990, -0.1414, "transfer", 15, aliases=("whitehall",)),
    _loc("Trafalgar Square", 51.5080, -0.1281, "landmark", 45, aliases=("trafalgar square",)),
    _loc("National Gallery", 51.5089, -0.1283, "museum", 60, aliases=("national gallery",)),
    _loc("Fortnum & Mason", 51.5097, -0.1396, "shopping", 60, aliases=("fortnum", "fortnum & mason")),
    _loc("Piccadilly Arcade", 51.5080, -0.1380, "shopping", 40, aliases=("piccadilly",)),
    _loc("British Museum", 51.5194, -0.1269, "museum", 90, aliases=("british museum",)),
    _loc("St Martin-in-the-Fields", 51.5091, -0.1269, "landmark", 30, aliases=("st martin",)),
    _loc("St Paul's Cathedral", 51.5138, -0.0984, "landmark", 60, aliases=("st paul", "st paul's cathedral")),
    _loc("Sky Garden", 51.5107, -0.0830, "landmark", 45, aliases=("sky garden",)),
    _loc("Leadenhall Market", 51.5123, -0.0830, "walk", 40, aliases=("leadenhall", )),
    _loc("Borough Market", 51.5054, -0.0911, "food", 60, aliases=("borough market",)),
    _loc("The Shard", 51.5045, -0.0865, "landmark", 40, aliases=("the shard", "shard")),
    _loc("Tower Bridge", 51.5055, -0.0754, "landmark", 45, aliases=("tower bridge",)),
    _loc("Tower of London", 51.5081, -0.0759, "landmark", 60, aliases=("tower of london",)),
    _loc("Camden Market", 51.5416, -0.1460, "experience", 90, aliases=("camden", "camden market")),
    _loc("Seabird Rooftop", 51.5084, -0.1002, "food", 90, aliases=("seabird",)),
    _loc("Hawksmoor Borough", 51.5059, -0.0919, "food", 90, aliases=("hawksmoor", "hawksmoor borough")),
    _loc("OXO Tower", 51.5076, -0.1117, "food", 90, aliases=("oxo tower", "oxo tower bar", "oxo tower brasserie")),
    _loc("Oblix at The Shard", 51.5045, -0.0865, "food", 120, aliases=("oblix", "oblix the shard")),
    _loc("The Gentlemen Baristas Bankside", 51.5066, -0.0999, "food", 30, aliases=("gentlemen baristas",)),
    _loc("Somerset House", 51.5113, -0.1163, "landmark", 45, aliases=("somerset house",)),
]


ESSENTIALS_BY_THEME: Dict[str, List[str]] = {
    "arrival": [
        "citizenM London Bankside",
        "South Bank Promenade",
        "London Eye",
        "Carnaby Street",
        "Regent Street",
        "Covent Garden",
    ],
    "mayfair": [
        "Soho",
        "Liberty London",
        "New Bond Street",
        "Fortnum & Mason",
        "Buckingham Palace",
        "St James's Park",
        "Hyde Park",
    ],
    "city": [
        "South Bank Promenade",
        "London Eye",
        "Borough Market",
        "Westminster Pier",
        "Trafalgar Square",
        "Piccadilly Arcade",
        "Fortnum & Mason",
        "Harrods",
        "Tower Bridge",
        "Oblix at The Shard",
    ],
    "departure": [
        "The Gentlemen Baristas Bankside",
        "citizenM London Bankside",
    ],
}

HOTEL_NAME = "citizenM London Bankside"

TRANSIT_OVERRIDES: Dict[Tuple[str, str], int] = {
    ("Heathrow Airport", "Paddington Station"): 15,
    ("Paddington Station", HOTEL_NAME): 30,
    (HOTEL_NAME, "Paddington Station"): 30,
    ("Paddington Station", "Heathrow Airport"): 15,
    (HOTEL_NAME, "Heathrow Airport"): 75,
    ("Heathrow Airport", HOTEL_NAME): 75,
}


# ---------------------------------------------------------------------------
# Utilitaires
# ---------------------------------------------------------------------------


def normalize_text(text: str) -> str:
    ascii_text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    ascii_text = re.sub(r"[^a-z0-9 ]", " ", ascii_text.lower())
    return re.sub(r"\s+", " ", ascii_text).strip()


ALIAS_INDEX: Dict[str, str] = {}
LOCATION_MAP: Dict[str, LocationInfo] = {}
for loc in LOCATION_LIBRARY:
    LOCATION_MAP[loc.name] = loc
    for alias in {loc.name.lower(), *map(str.lower, loc.aliases)}:
        ALIAS_INDEX[normalize_text(alias)] = loc.name


EXCLUDED_CATEGORIES = {"museum"}
UNWANTED_KEYWORDS = ("cathedral", "church", "abbey")

LOCATION_DESCRIPTIONS: Dict[str, str] = {
    "citizenM London Bankside": "Check-in et petite pause pour poser les valises.",
    "South Bank Promenade": "Flane le long de la Tamise, ambiance street art et food stalls.",
    "London Eye": "Vue panoramique douce pour se mettre dans l'ambiance.",
    "Carnaby Street": "Rue pop aux illuminations iconiques, parfaite en debut de soiree.",
    "Regent Street": "Arches lumineuses et vitrines de Noel pour une balade tranquille.",
    "Covent Garden": "Sapin geant, stands gourmands et atmosphere musicale.",
    "Soho": "Quartier vivant, cafes cosy et ruelles creatives.",
    "Liberty London": "Grand magasin a colombages, selection mode et deco arty.",
    "New Bond Street": "Vitrines de luxe et decors scintillants.",
    "Fortnum & Mason": "Maison de the historique, parfait pour une pause sucrerie.",
    "Buckingham Palace": "Cliche photo devant la residence royale, ambiance carte postale.",
    "St James's Park": "Parc apaisant pour respirer entre deux balades urbaines.",
    "Hyde Park": "Grand bol d'air, Winter Wonderland juste a cote.",
    "Hyde Park Winter Wonderland": "Attractions festives, vin chaud et grande roue.",
    "Trafalgar Square": "Esplanade animee, musiciens et illuminations historiques.",
    "Piccadilly Arcade": "Passage couvert elegant juste a cote de Piccadilly Circus.",
    "Harrods": "Galeries gourmandes et vitrines legendaires de Noel.",
    "Borough Market": "Marche gourmet aux stands de Noel pour grignoter local.",
    "The Shard": "Gratte-ciel iconique offrant un panorama grandiose.",
    "Oblix at The Shard": "Diner panoramique avec vue sur toute la ville.",
    "Tower Bridge": "Balade au bord de la Tamise avec iconique pont victorien.",
    "Mayfair Arcades": "Arcanes elegantes, idees cadeaux haut de gamme.",
    "Sketch Mayfair": "Pause healthy ou gourmandise dans un cadre arty.",
    "Bavarian Village": "Chalets en bois, biere chaude et musique live.",
    "Alpen Village": "Chemin gourmand plus calme pour se poser.",
}

TRANSFER_NOTES: Dict[Tuple[str, str], str] = {
    ("Heathrow Airport", "citizenM London Bankside"): "Heathrow Express jusqu'a Paddington puis taxi/metro (75 min).",
    ("citizenM London Bankside", "Harrods"): "Metro Jubilee puis Piccadilly Line (30 min environ).",
    ("citizenM London Bankside", "Hyde Park"): "Metro Jubilee -> Green Park puis promenade (25 min).",
    ("citizenM London Bankside", "Oblix at The Shard"): "Marche 10 min jusqu'au Shard, profiter des elevateurs panorama.",
}

PREFERRED_ORDER: Dict[str, List[str]] = {
    "arrival": ["Carnaby Street", "Regent Street", "Oxford Street", "Oxford Circus", "Covent Garden", "South Bank Promenade"],
    "mayfair": [
        "Soho",
        "Liberty London",
        "New Bond Street",
        "Fortnum & Mason",
        "Buckingham Palace",
        "St James's Park",
        "Hyde Park",
        "Hyde Park Winter Wonderland",
        "Bavarian Village",
        "Alpen Village",
    ],
    "city": [
        "South Bank Promenade",
        "London Eye",
        "Borough Market",
        "Westminster Pier",
        "Trafalgar Square",
        "Piccadilly Arcade",
        "Fortnum & Mason",
        "Harrods",
        "Tower Bridge",
        "The Shard",
        "Oblix at The Shard",
    ],
    "departure": ["The Gentlemen Baristas Bankside"],
}


def clean_location_list(candidates: Sequence[str]) -> List[str]:
    cleaned: List[str] = []
    for name in candidates:
        info = LOCATION_MAP.get(name)
        if not info:
            continue
        if info.category in EXCLUDED_CATEGORIES:
            continue
        if any(keyword in name.lower() for keyword in UNWANTED_KEYWORDS):
            continue
        cleaned.append(name)
    return cleaned

def parse_docx(path: Path) -> List[DaySection]:
    """Analyse le document Word et renvoie la structure journaliere."""
    if not path.exists():
        raise FileNotFoundError(path)

    with ZipFile(path) as doczip:
        xml = doczip.read("word/document.xml")

    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    paragraphs: List[str] = []
    for p in root.findall(".//w:p", ns):
        texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
        if not texts:
            paragraphs.append("")
        else:
            paragraphs.append("".join(texts))

    days: List[DaySection] = []
    current_title: Optional[str] = None
    buffer: List[str] = []

    for idx, para in enumerate(paragraphs):
        if any(para.startswith(marker) for marker in DAY_MARKERS):
            if current_title is not None:
                day_index = len(days)
                days.append(
                    DaySection(
                        index=day_index,
                        title=current_title.strip(),
                        theme=infer_theme(current_title),
                        lines=[line.strip() for line in buffer if line.strip()],
                        timeline=extract_timeline(buffer),
                    )
                )
            current_title = para.strip()
            buffer = []
        elif current_title is not None:
            buffer.append(para)

    if current_title is not None:
        day_index = len(days)
        days.append(
            DaySection(
                index=day_index,
                title=current_title.strip(),
                theme=infer_theme(current_title),
                lines=[line.strip() for line in buffer if line.strip()],
                timeline=extract_timeline(buffer),
            )
        )

    return days


def extract_timeline(lines: Iterable[str]) -> List[TimelineItem]:
    """Recherche des blocs (heure, activite, details) dans le tableau existant."""
    items: List[TimelineItem] = []
    clean_lines = [line.strip().replace("–", "-").replace("—", "-") for line in lines if line.strip()]
    i = 0
    while i < len(clean_lines):
        current = clean_lines[i]
        match = re.match(r"^(\d{1,2}h\d{2}(?:-\d{1,2}h\d{2})?)", current)
        if match:
            time = match.group(1)
            activity = clean_lines[i + 1] if i + 1 < len(clean_lines) else ""
            details = clean_lines[i + 2] if i + 2 < len(clean_lines) else ""
            items.append(TimelineItem(time=time, activity=activity, details=details))
            i += 3
        else:
            i += 1
    return items


def infer_theme(title: str) -> str:
    norm = normalize_text(title)
    if "depart" in norm or "dernier" in norm or "mercredi" in norm:
        return "departure"
    if "panorama" in norm or "trafalgar" in norm or "mardi" in norm:
        return "city"
    if "mayfair" in norm or "hyde" in norm or "lundi" in norm:
        return "mayfair"
    return "arrival"


def find_locations_in_text(lines: Iterable[str]) -> List[str]:
    found: List[str] = []
    for line in lines:
        norm_line = normalize_text(line)
        for alias, name in ALIAS_INDEX.items():
            if alias and alias in norm_line and name not in found:
                found.append(name)
    return found


def get_day_time_bounds(day: DaySection) -> Tuple[int, int]:
    """Retourne (start_min, end_min)."""
    start = math.inf
    end = -math.inf
    for item in day.timeline:
        start_min, end_min = parse_time_range(item.time)
        start = min(start, start_min)
        end = max(end, end_min)
    if start is math.inf:
        # Valeurs par defaut selon la thematique
        presets = {
            "arrival": (15 * 60, 23 * 60),
            "mayfair": (9 * 60 + 30, 22 * 60),
            "city": (9 * 60, 21 * 60 + 30),
            "departure": (7 * 60, 12 * 60),
        }
        return presets.get(day.theme, (9 * 60, 21 * 60))
    start = int(start)
    end = int(end)
    if day.theme == "city" and end < 22 * 60 + 30:
        end = 22 * 60 + 30
    return start, min(end, 23 * 60 + 30)


def parse_time_range(time_str: str) -> Tuple[int, int]:
    """Convertit une chaine 'HHhMM-HHhMM' en minutes absolues."""
    normalized = time_str.replace("–", "-").replace("—", "-")
    matches = re.findall(r"(\d{1,2})h(\d{2})", normalized)
    if not matches:
        raise ValueError(f"Format dheure inattendu: {time_str}")

    start_hour, start_min = map(int, matches[0])
    start_total = start_hour * 60 + start_min

    if len(matches) > 1:
        end_hour, end_min = map(int, matches[1])
        end_total = end_hour * 60 + end_min
    else:
        # par defaut, ajoute 60 min
        end_total = start_total + 60
    return start_total, end_total


DEDUP_WHITELIST = {
    HOTEL_NAME,
    "Heathrow Airport",
    "Paddington Station",
    "Southwark Station",
    "South Bank Promenade",
    "Bankside Pier",
    "Westminster Pier",
}


def merge_duplicate_locations(days: List[DaySection]) -> None:
    index_by_loc: Dict[str, List[int]] = {}
    for day in days:
        for name in day.locations:
            index_by_loc.setdefault(name, []).append(day.index)

    centroids = {day.index: compute_centroid(day.locations or [HOTEL_NAME]) for day in days}

    for name, indices in index_by_loc.items():
        if len(indices) <= 1:
            continue
        loc_info = LOCATION_MAP.get(name)
        if name not in DEDUP_WHITELIST and (not loc_info or loc_info.category not in {"transport", "hotel"}):
            continue
        best_index = min(indices, key=lambda idx: distance_between(name, centroids[idx]))
        for idx in indices:
            if idx == best_index:
                continue
            day = days[idx]
            if name == HOTEL_NAME and day.theme == "arrival":
                continue
            if name in day.locations:
                day.locations.remove(name)
                day.removed_duplicates.append(name)


def compute_centroid(location_names: Sequence[str]) -> Tuple[float, float]:
    lat_sum = 0.0
    lon_sum = 0.0
    count = 0
    for name in location_names:
        info = LOCATION_MAP.get(name)
        if not info:
            continue
        lat_sum += info.lat
        lon_sum += info.lon
        count += 1
    if count == 0:
        hotel = LOCATION_MAP[HOTEL_NAME]
        return hotel.lat, hotel.lon
    return lat_sum / count, lon_sum / count


def distance_between(location_name: str, point: Tuple[float, float]) -> float:
    loc = LOCATION_MAP.get(location_name)
    if not loc:
        return float("inf")
    return haversine_km((loc.lat, loc.lon), point)


def haversine_km(a: Tuple[float, float], b: Tuple[float, float]) -> float:
    lat1, lon1 = map(math.radians, a)
    lat2, lon2 = map(math.radians, b)
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    h = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
    return 6371.0 * (2 * math.asin(math.sqrt(h)))


def add_missing_essentials(days: List[DaySection]) -> None:
    used = {loc for day in days for loc in day.locations}
    for day in days:
        essentials = ESSENTIALS_BY_THEME.get(day.theme, [])
        centroid = compute_centroid(day.locations or [HOTEL_NAME])
        for essential in essentials:
            if essential in used:
                continue
            info = LOCATION_MAP.get(essential)
            if not info:
                continue
            if info.category in EXCLUDED_CATEGORIES:
                continue
            if any(keyword in essential.lower() for keyword in UNWANTED_KEYWORDS):
                continue
            used.add(essential)
            day.locations.append(essential)
            day.added_essentials.append(essential)
        # Nettoyage puis tri par distance pour une sequence coherente
        unique_locations = list(dict.fromkeys(clean_location_list(day.locations)))
        day.locations = sorted(unique_locations, key=lambda name: distance_between(name, centroid))


def order_locations(start: str, visits: Sequence[str], end: str) -> List[str]:
    remaining = [loc for loc in visits if loc not in {start, end}]
    ordered: List[str] = []
    current = start
    while remaining:
        next_loc = min(remaining, key=lambda name: distance_between_pair(current, name))
        ordered.append(next_loc)
        remaining.remove(next_loc)
        current = next_loc
    if end and (not ordered or ordered[-1] != end):
        ordered.append(end)
    return ordered


def distance_between_pair(a_name: str, b_name: str) -> float:
    if a_name == b_name:
        return 0.0
    a = LOCATION_MAP.get(a_name)
    b = LOCATION_MAP.get(b_name)
    if not a or not b:
        return 999.0
    return haversine_km((a.lat, a.lon), (b.lat, b.lon))


def estimate_travel_minutes(a_name: str, b_name: str) -> int:
    if a_name == b_name:
        return 0
    if (a_name, b_name) in TRANSIT_OVERRIDES:
        return TRANSIT_OVERRIDES[(a_name, b_name)]
    if (b_name, a_name) in TRANSIT_OVERRIDES:
        return TRANSIT_OVERRIDES[(b_name, a_name)]

    distance = distance_between_pair(a_name, b_name)
    loc_a = LOCATION_MAP.get(a_name)
    loc_b = LOCATION_MAP.get(b_name)

    if not loc_a or not loc_b:
        return int(distance / 4.5 * 60) + 5

    base_speed_kmh = 3.8 if all(cat in {"walk", "landmark", "shopping", "park"} for cat in (loc_a.category, loc_b.category)) else 5.0
    minutes = int((distance / max(base_speed_kmh, 0.1)) * 60)
    return max(minutes, 8)


def describe_location(name: str) -> str:
    if name in LOCATION_DESCRIPTIONS:
        return LOCATION_DESCRIPTIONS[name]
    info = LOCATION_MAP.get(name)
    if not info:
        return ""
    if info.category == "food":
        return "Pause gourmande en douceur."
    if info.category == "walk":
        return "Balade libre pour profiter du quartier."
    if info.category == "landmark":
        return "Point de vue iconique a decouvrir tranquillement."
    return info.category.capitalize()


def friendly_transfer_detail(start: str, end: str, duration: int) -> str:
    note = TRANSFER_NOTES.get((start, end)) or TRANSFER_NOTES.get((end, start))
    if note:
        return note
    if duration <= 15:
        return "Petite marche ou trajet rapide (<= 15 min)."
    if duration <= 30:
        return "Deplacement fluide (~30 min), prends ton temps."
    return "Prevoir ce trajet sans stress, musique ou podcast en route."


def build_day_schedule(day: DaySection) -> List[Segment]:
    start_time, end_time = get_day_time_bounds(day)
    current_time = start_time

    segments: List[Segment] = []

    poi_locations: List[str] = []
    for loc in day.locations:
        info = LOCATION_MAP.get(loc)
        if not info:
            continue
        if info.category in {"transport", "hotel"}:
            continue
        poi_locations.append(loc)

    if "Oblix at The Shard" in poi_locations and "The Shard" in poi_locations:
        poi_locations = [loc for loc in poi_locations if loc != "The Shard"]

    previous = HOTEL_NAME

    if day.theme == "arrival":
        arrival = LOCATION_MAP["Heathrow Airport"]
        segments.append(
            Segment(
                start=current_time,
                end=current_time + arrival.default_duration,
                title="Arrivee a Heathrow",
                location="Heathrow Airport",
                details="Formalites et recuperation des valises.",
                segment_type="transport",
            )
        )
        current_time += arrival.default_duration

        travel_duration = estimate_travel_minutes("Heathrow Airport", HOTEL_NAME)
        segments.append(
            Segment(
                start=current_time,
                end=current_time + travel_duration,
                title="Vers l'hotel",
                location="Heathrow Airport -> citizenM London Bankside",
                details=friendly_transfer_detail("Heathrow Airport", HOTEL_NAME, travel_duration),
                segment_type="transfer",
            )
        )
        current_time += travel_duration

        hotel_info = LOCATION_MAP[HOTEL_NAME]
        checkin_duration = max(45, hotel_info.default_duration)
        segments.append(
            Segment(
                start=current_time,
                end=current_time + checkin_duration,
                title="Installation a l'hotel",
                location=HOTEL_NAME,
                details=describe_location(HOTEL_NAME),
                segment_type="visit",
            )
        )
        current_time += checkin_duration
        previous = HOTEL_NAME

    else:
        previous = HOTEL_NAME

    visits_order: List[str] = []
    preferred = PREFERRED_ORDER.get(day.theme, [])
    for name in preferred:
        if name in poi_locations and name not in visits_order:
            visits_order.append(name)
    for name in poi_locations:
        if name not in visits_order:
            visits_order.append(name)
    if day.theme == "city" and "Oblix at The Shard" not in visits_order:
        visits_order.append("Oblix at The Shard")

    for next_loc in visits_order:
        loc_info = LOCATION_MAP.get(next_loc)
        if not loc_info:
            continue

        travel_duration = estimate_travel_minutes(previous, next_loc)
        if current_time + travel_duration > end_time:
            break
        if travel_duration:
            segments.append(
                Segment(
                    start=current_time,
                    end=current_time + travel_duration,
                    title=f"Trajet vers {next_loc}",
                    location=f"{previous} -> {next_loc}",
                    details=friendly_transfer_detail(previous, next_loc, travel_duration),
                    segment_type="transfer",
                )
            )
            current_time += travel_duration

        visit_duration = loc_info.default_duration
        if loc_info.category == "food":
            visit_duration = max(75, visit_duration)
        else:
            visit_duration = max(45, visit_duration)

        if current_time + visit_duration > end_time:
            visit_duration = max(30, end_time - current_time)
        if visit_duration <= 0:
            continue

        segments.append(
            Segment(
                start=current_time,
                end=current_time + visit_duration,
                title=next_loc,
                location=next_loc,
                details=describe_location(next_loc),
                segment_type="meal" if loc_info.category == "food" else "visit",
            )
        )
        current_time += visit_duration
        previous = next_loc

    if day.theme != "departure" and previous != HOTEL_NAME:
        travel_duration = estimate_travel_minutes(previous, HOTEL_NAME)
        segments.append(
            Segment(
                start=current_time,
                end=current_time + travel_duration,
                title="Retour a l'hotel",
                location=f"{previous} -> {HOTEL_NAME}",
                details=friendly_transfer_detail(previous, HOTEL_NAME, travel_duration),
                segment_type="transfer",
            )
        )
        current_time += travel_duration
        if current_time < end_time:
            segments.append(
                Segment(
                    start=current_time,
                    end=end_time,
                    title="Fin de soiree detendue",
                    location=HOTEL_NAME,
                    details="Temps libre pour un verre ou repos.",
                    segment_type="buffer",
                )
            )

    if day.theme == "departure":
        if previous != HOTEL_NAME:
            travel_duration = estimate_travel_minutes(previous, HOTEL_NAME)
            segments.append(
                Segment(
                    start=current_time,
                    end=current_time + travel_duration,
                    title="Retour a l'hotel",
                    location=f"{previous} -> {HOTEL_NAME}",
                    details=friendly_transfer_detail(previous, HOTEL_NAME, travel_duration),
                    segment_type="transfer",
                )
            )
            current_time += travel_duration

        transfer = estimate_travel_minutes(HOTEL_NAME, "Heathrow Airport")
        segments.append(
            Segment(
                start=current_time,
                end=current_time + transfer,
                title="Transfert vers Heathrow",
                location=f"{HOTEL_NAME} -> Heathrow Airport",
                details=friendly_transfer_detail(HOTEL_NAME, "Heathrow Airport", transfer),
                segment_type="transfer",
            )
        )
        current_time += transfer
        segments.append(
            Segment(
                start=current_time,
                end=current_time + 45,
                title="Formalites de depart",
                location="Heathrow Airport",
                details="Enregistrement et controle securite.",
                segment_type="transport",
            )
        )

    return segments


def minutes_to_label(minutes: int) -> str:
    hour = minutes // 60
    minute = minutes % 60
    return f"{hour:02d}h{minute:02d}"


def make_time_range(start: int, end: int) -> str:
    return f"{minutes_to_label(start)}-{minutes_to_label(end)}"


def ensure_directory(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def create_daily_map(day: DaySection, segments: List[Segment], output_dir: Path) -> Path:
    start_name = "Heathrow Airport" if day.theme == "arrival" else HOTEL_NAME
    end_name = "Heathrow Airport" if day.theme == "departure" else HOTEL_NAME

    start_info = LOCATION_MAP.get(start_name)
    end_info = LOCATION_MAP.get(end_name)

    visits = [s for s in segments if s.segment_type in {"visit", "meal"} and s.location in LOCATION_MAP]
    if not start_info or (not visits and not end_info):
        return output_dir / f"day_{day.index + 1}_no_map.html"

    route_sequence: List[str] = []
    if start_info:
        route_sequence.append(start_name)
    for seg in visits:
        if seg.location not in route_sequence:
            route_sequence.append(seg.location)
    if end_info and (not route_sequence or route_sequence[-1] != end_name):
        route_sequence.append(end_name)

    coords = [(LOCATION_MAP[name].lat, LOCATION_MAP[name].lon) for name in route_sequence if name in LOCATION_MAP]
    if not coords:
        return output_dir / f"day_{day.index + 1}_no_map.html"

    avg_lat = sum(lat for lat, _ in coords) / len(coords)
    avg_lon = sum(lon for _, lon in coords) / len(coords)

    fmap = folium.Map(location=[avg_lat, avg_lon], zoom_start=14)

    # Start marker
    if start_info:
        folium.Marker(
            location=[start_info.lat, start_info.lon],
            tooltip=f"Depart : {start_name}",
            popup=f"0. Depart - {start_name}",
            icon=folium.Icon(color="red", icon="home"),
        ).add_to(fmap)

    route_points = [[start_info.lat, start_info.lon]] if start_info else []

    for idx, seg in enumerate(visits, start=1):
        loc = LOCATION_MAP.get(seg.location)
        if not loc:
            continue
        popup = f"{idx}. {seg.title} ({minutes_to_label(seg.start)} - {minutes_to_label(seg.end)})"
        details = seg.details or ""
        if details:
            popup = f"{popup}<br>{details}"
        folium.Marker(
            location=[loc.lat, loc.lon],
            tooltip=seg.title,
            popup=popup,
            icon=folium.Icon(color="green" if seg.segment_type == "meal" else "blue", icon="info-sign"),
        ).add_to(fmap)
        route_points.append([loc.lat, loc.lon])

    if end_info:
        folium.Marker(
            location=[end_info.lat, end_info.lon],
            tooltip=f"Retour : {end_name}",
            popup=f"Fin - {end_name}",
            icon=folium.Icon(color="purple", icon="flag"),
        ).add_to(fmap)
        route_points.append([end_info.lat, end_info.lon])

    if len(route_points) >= 2:
        folium.PolyLine(route_points, color="#FF6F61", weight=4, opacity=0.8, tooltip="Parcours du jour").add_to(fmap)

    legend_html = """
    <div style="
        position: fixed;
        bottom: 20px;
        left: 20px;
        z-index: 9999;
        background-color: white;
        padding: 10px 14px;
        border: 1px solid #ccc;
        box-shadow: 0 0 6px rgba(0,0,0,0.2);
        font-size: 12px;
        line-height: 1.4;
    ">
        <b>Legende</b><br>
        <span style="color:#d9534f">&#9679;</span> Depart<br>
        <span style="color:#337ab7">&#9679;</span> Balade / visite<br>
        <span style="color:#5cb85c">&#9679;</span> Pause gourmande<br>
        <span style="color:#6f42c1">&#9679;</span> Retour hotel / aeroport<br>
        <span style="color:#FF6F61">&#8213;</span> Trajet conseille
    </div>
    """
    legend = MacroElement()
    legend._template = Template(
        "{% macro html(this, kwargs) %}" + legend_html + "{% endmacro %}"
    )
    fmap.get_root().add_child(legend)

    map_path = output_dir / f"day_{day.index + 1}_{slugify(day.title)}.html"
    fmap.save(str(map_path))
    return map_path


def slugify(text: str) -> str:
    normalized = normalize_text(text)
    return normalized.replace(" ", "_") or "jour"


def to_windows_uri(path: Path) -> str:
    resolved = path.resolve()
    raw = str(resolved)
    if raw.startswith("/mnt/") and len(raw) > 6:
        drive_letter = raw[5].upper()
        windows_path = raw[7:].replace("/", "/")
        return f"file:///{drive_letter}:/{quote(windows_path)}"
    return resolved.as_uri()


def add_hyperlink(paragraph, text: str, url_path: Path) -> None:
    """Ajoute un lien cliquable dans un paragraphe python-docx."""
    href = to_windows_uri(url_path)
    part = paragraph.part
    r_id = part.relate_to(href, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_document(output_path: Path, days: List[DaySection], day_segments: Dict[int, List[Segment]], map_paths: Dict[int, Path]) -> None:
    doc = Document()
    styles = doc.styles
    if "Itinerary Heading" not in styles:
        heading_style = styles.add_style("Itinerary Heading", WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = styles["Heading 2"]

    doc.add_heading("Itineraire optimise - Sejour a Londres", level=1)
    intro = doc.add_paragraph(
        "Version generee automatiquement : regroupement malin des activites par quartiers et rythme detendu pour profiter sans stress."
    )
    intro.runs[0].italic = True
    note = doc.add_paragraph("Cartes interactives : connexion internet requise pour afficher les fonds OpenStreetMap.")
    if note.runs:
        note.runs[0].italic = True

    for day in days:
        doc.add_paragraph(day.title, style="Itinerary Heading")
        summary = doc.add_paragraph()
        summary.add_run("Optimisations : ")
        changes: List[str] = []
        if day.removed_duplicates:
            removed = ", ".join(day.removed_duplicates)
            changes.append(f"activites regroupees ailleurs ({removed})")
        if day.added_essentials:
            added = ", ".join(day.added_essentials)
            changes.append(f"incontournables ajoutes ({added})")
        if not changes:
            changes.append("reorganisation par proximite geographique")
        summary.add_run("; ".join(changes))

        segments = day_segments.get(day.index, [])
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Horaire"
        hdr[1].text = "Moment"
        hdr[2].text = "Ambiance"

        for seg in segments:
            row = table.add_row().cells
            row[0].text = make_time_range(seg.start, seg.end)
            row[1].text = seg.title
            row[2].text = seg.details

        doc.add_paragraph("")
        map_path = map_paths.get(day.index)
        if map_path and map_path.exists():
            para = doc.add_paragraph("Carte interactive : ")
            add_hyperlink(para, map_path.name, map_path.resolve())
        doc.add_paragraph("")  # espace

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Optimise l'itineraire et genere cartes/Word.")
    parser.add_argument("--output-dir", default="optimized_london", help="Repertoire de sortie (defaut: optimized_london).")
    parser.add_argument("--skip-document", action="store_true", help="Ne pas regenerer le document Word (utile si modifie manuellement).")
    parser.add_argument("--skip-maps", action="store_true", help="Ne pas regenerer les cartes HTML.")
    args = parser.parse_args()

    input_path = Path("London - mise \u00e0 jour.docx")
    output_dir = ensure_directory(Path(args.output_dir))
    map_dir = ensure_directory(output_dir / "maps") if not args.skip_maps else output_dir / "maps"
    output_doc = output_dir / "London_itinerary_optimise.docx"

    days = parse_docx(input_path)
    for day in days:
        located = find_locations_in_text([day.title, *day.lines])
        day.locations = clean_location_list(located)

    merge_duplicate_locations(days)
    add_missing_essentials(days)

    for day in days:
        if "Oblix at The Shard" in day.locations and "The Shard" in day.locations:
            day.locations = [loc for loc in day.locations if loc != "The Shard"]
        if day.theme == "city":
            day.locations = [loc for loc in day.locations if loc != "Hyde Park"]

    day_segments: Dict[int, List[Segment]] = {}
    map_paths: Dict[int, Path] = {}

    for day in days:
        segments = build_day_schedule(day)
        day_segments[day.index] = segments
        if not args.skip_maps:
            map_paths[day.index] = create_daily_map(day, segments, map_dir)

    if not args.skip_document:
        build_document(output_doc, days, day_segments, map_paths)

    result_summary = {
        "output_document": None if args.skip_document else str(output_doc),
        "maps": {idx: str(path) for idx, path in map_paths.items()},
        "changes": {
            day.title: {
                "removed_duplicates": day.removed_duplicates,
                "added_essentials": day.added_essentials,
                "locations": day.locations,
            }
            for day in days
        },
    }

    with open(output_dir / "summary.json", "w", encoding="utf-8") as fh:
        json.dump(result_summary, fh, ensure_ascii=False, indent=2)

    print(json.dumps(result_summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
